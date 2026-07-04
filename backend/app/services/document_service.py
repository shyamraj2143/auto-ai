import hashlib
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

from app.core.config import settings
from app.services.groq_service import groq_service


@dataclass(frozen=True)
class DocumentExtraction:
    text: str
    metadata: dict[str, Any]


class DocumentService:
    async def save_and_extract(
        self,
        upload: UploadFile,
        user_id: str,
        *,
        max_upload_mb: int | None = None,
    ) -> tuple[str, DocumentExtraction]:
        filename = upload.filename or "document"
        extension = Path(filename).suffix.lower()
        if extension not in settings.ALLOWED_DOCUMENT_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Supported document formats are PDF, TXT, and DOCX.",
            )

        data = await upload.read()
        if not data:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="The uploaded file is empty.",
            )

        upload_limit_mb = max_upload_mb or settings.MAX_UPLOAD_MB
        max_bytes = upload_limit_mb * 1024 * 1024
        if len(data) > max_bytes:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File exceeds {upload_limit_mb} MB.",
            )

        extraction = self.extract_text(data, extension)
        if not extraction.text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="No readable text was found in the document.",
            )

        safe_name = self.safe_filename(filename)
        user_dir = Path(settings.UPLOAD_DIR) / user_id
        user_dir.mkdir(parents=True, exist_ok=True)
        digest = hashlib.sha256(data).hexdigest()[:18]
        stored_path = user_dir / f"{Path(safe_name).stem}-{digest}{extension}"
        stored_path.write_bytes(data)
        metadata = {
            **extraction.metadata,
            "original_filename": filename,
            "stored_filename": stored_path.name,
            "extension": extension,
            "content_type": upload.content_type or "application/octet-stream",
            "file_size": len(data),
            "sha256": digest,
        }
        return str(stored_path), DocumentExtraction(text=extraction.text, metadata=metadata)

    def extract_text(self, data: bytes, extension: str) -> DocumentExtraction:
        try:
            if extension == ".txt":
                return self._extract_txt(data)
            if extension == ".pdf":
                return self._extract_pdf(data)
            if extension == ".docx":
                return self._extract_docx(data)
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Could not read this {extension.upper().lstrip('.')} file. It may be encrypted or corrupted.",
            ) from exc
        return self._build_extraction("", parser="unknown")

    def _extract_txt(self, data: bytes) -> DocumentExtraction:
        text = ""
        for encoding in ("utf-8-sig", "utf-16", "latin-1"):
            try:
                text = data.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if not text:
            text = data.decode("utf-8", errors="replace")
        return self._build_extraction(text, parser="text")

    def _extract_pdf(self, data: bytes) -> DocumentExtraction:
        reader = PdfReader(BytesIO(data), strict=False)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Encrypted PDFs are not supported unless they can be opened without a password.",
                ) from exc

        pages: list[str] = []
        failed_pages: list[int] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                failed_pages.append(index)
                pages.append("")

        return self._build_extraction(
            "\n\n".join(pages),
            parser="pypdf",
            page_count=len(reader.pages),
            failed_pages=failed_pages,
        )

    def _extract_docx(self, data: bytes) -> DocumentExtraction:
        doc = DocxDocument(BytesIO(data))
        blocks: list[str] = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    blocks.append(" | ".join(values))
        return self._build_extraction("\n".join(blocks), parser="python-docx")

    def _build_extraction(self, text: str, **metadata: Any) -> DocumentExtraction:
        cleaned = self.normalize_text(text)
        return DocumentExtraction(
            text=cleaned,
            metadata={
                **metadata,
                "character_count": len(cleaned),
                "word_count": len(re.findall(r"\S+", cleaned)),
            },
        )

    def summarize(self, extracted_text: str, filename: str, *, provider: str | None = None) -> str:
        text = extracted_text[: settings.MAX_DOCUMENT_CONTEXT_CHARS]
        messages = [
            {
                "role": "system",
                "content": "You are Auto-AI. Produce concise, accurate document summaries.",
            },
            {
                "role": "user",
                "content": (
                    f"Summarize the document named {filename}. Include key points, decisions, "
                    f"risks, and action items when present.\n\n{text}"
                ),
            },
        ]
        provider_chain = self._provider_chain(provider)
        last_error: Exception | None = None

        for selected_provider in provider_chain:
            try:
                content, _, _ = groq_service.complete(messages, provider=selected_provider)
                return content
            except HTTPException as exc:
                last_error = exc
                continue

        if last_error:
            raise last_error
        return ""

    def document_context(self, documents: list[tuple[str, str]]) -> str:
        if not documents:
            return ""

        budget = settings.MAX_DOCUMENT_CONTEXT_CHARS
        chunks: list[str] = []
        for filename, text in documents:
            if budget <= 0:
                break
            excerpt = text[: min(len(text), budget)]
            budget -= len(excerpt)
            chunks.append(f"Document: {filename}\n{excerpt}")
        return "\n\n---\n\n".join(chunks)

    @staticmethod
    def safe_filename(filename: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", filename).strip("._")
        return cleaned or "document"

    @staticmethod
    def normalize_text(text: str) -> str:
        text = text.replace("\x00", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _provider_chain(provider: str | None) -> list[str]:
        preferred = (provider or settings.AI_PROVIDER).lower()
        chain = [preferred]
        if preferred != "groq":
            chain.append("groq")
        if preferred != "openai":
            chain.append("openai")
        return chain


document_service = DocumentService()
